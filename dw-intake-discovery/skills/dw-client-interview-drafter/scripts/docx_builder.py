#!/usr/bin/env python3
"""
docx_builder.py — generic .docx generator for dw-client-interview-drafter.

INPUT SCHEMA (build a Python dict and pass via build()):

    {
        "mode": "long-form-initial" | "short-form-initial"
              | "follow-up-post-discovery" | "follow-up-pre-plea"
              | "follow-up-pre-trial" | "follow-up-post-motion-ruling",
        "case_context": {
            "client_name": "Last, First",
            "dob": "M/D/YYYY",
            "arrest_date": "M/D/YYYY",
            "first_appearance": "M/D/YYYY",
            "court": "Orleans CDC / Sec. M4",
            "magistrate_num": "...",
            "folder_num": "...",
            "ccn": "...",
            "jail_facility": "OJC / OPP",
            "interpreter": "No / Yes (lang)",
            "charges": [
                {"statute": "La. R.S. 14:65", "desc": "Simple Robbery", "class": "Felony - VIOLENT"},
                ...
            ],
            "priors": [
                {"date": "2019", "case": "544303L", "charge": "...", "disposition": "..."},
                ...
            ],
            "risk_callouts": ["PSA Risk Level 4 ...", "Habitual exposure ..."],
            "trial_date": "M/D/YYYY (pre-trial mode only)",
            "motion_ruled": "Suppression (pre-motion mode only)",
            "ruling_outcome": "Denied / Granted / ...",
            "plea_offer_summary": "for pre-plea mode only",
        },
        "modules": ["firearms", "violent-felony", "drug-offense"],  # charge module slugs to inject
        "output_path": "/abs/path/to/output.docx",
        "skill_root": "/abs/path/to/dw-client-interview-drafter",  # to find references
    }

The script reads the matching reference .md and renders it into a structured .docx.
"""

import os
import re
import sys
import json
from datetime import date

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    sys.stderr.write("python-docx required: pip install python-docx --break-system-packages\n")
    sys.exit(1)

MODE_TEMPLATES = {
    "long-form-initial":           "references/long-form-initial.md",
    "short-form-initial":          "references/short-form-initial.md",
    "follow-up-post-discovery":    "references/follow-up-post-discovery.md",
    "follow-up-pre-plea":          "references/follow-up-pre-plea.md",
    "follow-up-pre-trial":         "references/follow-up-pre-trial.md",
    "follow-up-post-motion-ruling":"references/follow-up-post-motion-ruling.md",
}

MODE_TITLES = {
    "long-form-initial":            "INITIAL CLIENT INTERVIEW - QUESTION SHEET",
    "short-form-initial":           "INITIAL CLIENT TRIAGE - JAIL VISIT SHEET",
    "follow-up-post-discovery":     "CLIENT FOLLOW-UP INTERVIEW - POST-DISCOVERY REVIEW",
    "follow-up-pre-plea":           "CLIENT FOLLOW-UP INTERVIEW - PRE-PLEA CONSULTATION",
    "follow-up-pre-trial":          "CLIENT FOLLOW-UP INTERVIEW - PRE-TRIAL READINESS",
    "follow-up-post-motion-ruling": "CLIENT FOLLOW-UP INTERVIEW - POST-RULING DEBRIEF",
}

# ---------- style helpers ----------
def setup_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)
    return doc

def p(doc, text, *, bold=False, italic=False, size=10.5, align=None, color=None, indent=None, after=4, before=0):
    para = doc.add_paragraph()
    if align is not None: para.alignment = align
    para.paragraph_format.space_after = Pt(after)
    para.paragraph_format.space_before = Pt(before)
    if indent is not None: para.paragraph_format.left_indent = Inches(indent)
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    return para

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
DARK_RED = RGBColor(0xB0, 0x00, 0x00)
GREY = RGBColor(0x60, 0x60, 0x60)

def h1(doc, text):  p(doc, text, bold=True, size=14, color=NAVY, after=4, before=8)
def h2(doc, text):  p(doc, text, bold=True, size=12, color=NAVY, after=4, before=10)
def h3(doc, text):  p(doc, text, bold=True, italic=True, size=10.5, after=2, before=6)
def q(doc, text):   p(doc, "*  " + text, size=10.5, indent=0.25, after=8)
def note(doc, text):p(doc, text, italic=True, size=9.5, color=GREY, indent=0.25, after=4)
def warn(doc, text):p(doc, ">> " + text, bold=True, size=10, color=DARK_RED, indent=0.25, after=6)

# ---------- markdown parser (lightweight, tailored to our templates) ----------
def parse_template(doc, md_text):
    """Render the reference markdown into the doc. Recognized syntax:
       # H1, ## H2, ### H3, **Warning:** / **Guidance:** lines, bullet items '- ', horizontal rules ---."""
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped:
            i += 1
            continue
        if stripped == '---':
            i += 1
            continue
        # H1
        if stripped.startswith('# ') and not stripped.startswith('## '):
            h1(doc, stripped[2:].strip())
        # H2
        elif stripped.startswith('## ') and not stripped.startswith('### '):
            h2(doc, stripped[3:].strip())
        # H3
        elif stripped.startswith('### '):
            h3(doc, stripped[4:].strip())
        # Bold callouts (Warning, Guidance)
        elif stripped.startswith('**Warning:**'):
            warn(doc, stripped.replace('**Warning:**', '').strip())
        elif stripped.startswith('**Guidance:**'):
            note(doc, stripped.replace('**Guidance:**', '').strip())
        elif stripped.startswith('**Hard rule:**'):
            warn(doc, "HARD RULE: " + stripped.replace('**Hard rule:**', '').strip())
        # Bullets
        elif stripped.startswith('- '):
            q(doc, stripped[2:].strip())
        # Inline code line (signature placeholders, etc.)
        elif stripped.startswith('`') and stripped.endswith('`'):
            p(doc, stripped.strip('`'), size=10, italic=True, color=GREY, after=4, before=4)
        # Italic single line guidance "*[...]*" injection markers
        elif stripped.startswith('*[') and stripped.endswith(']*'):
            note(doc, stripped.strip('*'))
        # Plain paragraph
        else:
            p(doc, stripped, size=10.5, after=4)
        i += 1

# ---------- case header rendering ----------
def render_header(doc, mode, ctx):
    p(doc, "ATTORNEY WORK PRODUCT - PRIVILEGED & CONFIDENTIAL",
      bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=DARK_RED, after=2)
    p(doc, "Prepared by Cowork at the direction of counsel - Draft for attorney review",
      italic=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY, after=10)
    h1(doc, MODE_TITLES[mode])
    p(doc, f"Daniels & Washington | Criminal Defense | {ctx.get('court', '[Court]')}",
      italic=True, size=10, color=RGBColor(0x40, 0x40, 0x40), after=10)

    # Case header table
    rows = [
        ["Client:", ctx.get("client_name", ""),         "DOB:",              ctx.get("dob", "")],
        ["Arrest Date:", ctx.get("arrest_date", ""),    "First Appearance:", ctx.get("first_appearance", "")],
        ["Court / Section:", ctx.get("court", ""),      "Magistrate #:",     ctx.get("magistrate_num", "")],
        ["Folder #:", ctx.get("folder_num", ""),        "CCN #:",            ctx.get("ccn", "")],
        ["Jail Facility:", ctx.get("jail_facility", ""),"Interpreter:",      ctx.get("interpreter", "No")],
    ]
    t = doc.add_table(rows=len(rows), cols=4)
    t.style = 'Light Grid Accent 1'
    for r, row in enumerate(rows):
        for c, v in enumerate(row):
            cell = t.rows[r].cells[c]
            cell.text = v
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9.5)
                    if c % 2 == 0:
                        run.bold = True
    p(doc, "", after=4)

    # Mode-specific extra header
    if mode == "follow-up-pre-trial" and ctx.get("trial_date"):
        p(doc, f"Trial Date: {ctx['trial_date']}", bold=True, size=10, after=4)
    if mode == "follow-up-post-motion-ruling":
        if ctx.get("motion_ruled"):
            p(doc, f"Motion Ruled On: {ctx['motion_ruled']}", bold=True, size=10, after=2)
        if ctx.get("ruling_outcome"):
            p(doc, f"Ruling Outcome: {ctx['ruling_outcome']}", bold=True, size=10, after=4)
    if mode == "follow-up-pre-plea" and ctx.get("plea_offer_summary"):
        warn(doc, "Plea offer (attorney-completed before meeting):")
        p(doc, ctx["plea_offer_summary"], size=10, indent=0.25, after=6)

    # Current charges
    if ctx.get("charges"):
        h2(doc, "CURRENT CHARGES")
        ct = doc.add_table(rows=1, cols=3)
        ct.style = 'Light List Accent 1'
        hdr = ct.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Statute", "Description", "Class / Notes"
        for ch in ctx["charges"]:
            row = ct.add_row().cells
            row[0].text = ch.get("statute", "")
            row[1].text = ch.get("desc", "")
            row[2].text = ch.get("class", "")
        for row in ct.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9.5)
        p(doc, "", after=2)

    # Risk callouts
    for cb in ctx.get("risk_callouts", []):
        warn(doc, cb)

def render_priors_table(doc, priors):
    if not priors:
        return
    h3(doc, "Prior Record (from PSA / NCIC / rap sheet — confirm each)")
    t = doc.add_table(rows=1, cols=4)
    t.style = 'Light List Accent 1'
    hdr = t.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Year / Case #", "Charge", "Disposition", "Confirm / Notes"
    for pr in priors:
        row = t.add_row().cells
        row[0].text = f"{pr.get('date', '')} / {pr.get('case', '')}".strip(" /")
        row[1].text = pr.get("charge", "")
        row[2].text = pr.get("disposition", "")
        row[3].text = pr.get("notes", "")
    for row in t.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    p(doc, "", after=2)

def render_module(doc, skill_root, module_slug):
    """Inject a charge module from references/charge-modules/<slug>.md."""
    path = os.path.join(skill_root, "references", "charge-modules", f"{module_slug}.md")
    if not os.path.exists(path):
        warn(doc, f"Module not found: {module_slug}")
        return
    with open(path) as f:
        parse_template(doc, f.read())

def render_signoff(doc, mode):
    p(doc, "", after=8)
    p(doc, "- END OF SHEET -", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY, size=9, after=2)
    if mode == "long-form-initial":
        p(doc, "Interview date: ________   Attorney: ________   Location: ________",
          size=9, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    elif mode == "short-form-initial":
        p(doc, "Visit date: ________   Attorney: ________   Time in: ____   Time out: ____   Facility: ________",
          size=9, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    elif mode == "follow-up-pre-trial":
        p(doc, "Interview date: ________   Attorney: ________   Trial date: ________   Location: ________",
          size=9, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    elif mode == "follow-up-pre-plea":
        p(doc, "Consultation date: ________   Attorney: ________   Witness (if any): ________   Location: ________",
          size=9, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    elif mode == "follow-up-post-motion-ruling":
        p(doc, "Interview date: ________   Attorney: ________   Motion: ________   Location: ________",
          size=9, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    else:
        p(doc, "Interview date: ________   Attorney: ________   Location: ________",
          size=9, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)

# ---------- entrypoint ----------
def build(spec):
    mode = spec["mode"]
    ctx = spec.get("case_context", {})
    modules = spec.get("modules", [])
    output_path = spec["output_path"]
    skill_root = spec["skill_root"]

    if mode not in MODE_TEMPLATES:
        raise ValueError(f"Unknown mode: {mode}")

    doc = setup_doc()
    render_header(doc, mode, ctx)

    template_path = os.path.join(skill_root, MODE_TEMPLATES[mode])
    with open(template_path) as f:
        md = f.read()

    # Long-form: split at the injection marker, place charge modules between F and I
    if mode == "long-form-initial":
        marker = "## [CHARGE MODULE INJECTION POINT]"
        if marker in md:
            before, after = md.split(marker, 1)
        else:
            before, after = md, ""
        parse_template(doc, before)
        # priors table goes inside Section J — but we render an early priors block for the attorney here
        # (Section J body still asks the verification questions)
        for slug in modules:
            render_module(doc, skill_root, slug)
        # render the priors table near the criminal history section
        # The simplest approach: inject priors right before "## SECTION J"
        if "## SECTION J" in after:
            j_before, j_after = after.split("## SECTION J", 1)
            parse_template(doc, j_before)
            h1(doc, "J.  CRIMINAL HISTORY VERIFICATION")
            render_priors_table(doc, ctx.get("priors", []))
            parse_template(doc, j_after.split("\n", 1)[1] if "\n" in j_after else "")
        else:
            parse_template(doc, after)
    else:
        parse_template(doc, md)
        # For follow-ups, modules are still useful if specified
        for slug in modules:
            render_module(doc, skill_root, slug)

    render_signoff(doc, mode)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return {
        "output_path": output_path,
        "mode": mode,
        "modules_injected": modules,
        "size_bytes": os.path.getsize(output_path),
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
    }

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("usage: docx_builder.py <spec.json>")
        sys.exit(2)
    with open(sys.argv[1]) as f:
        spec = json.load(f)
    result = build(spec)
    print(json.dumps(result, indent=2))
