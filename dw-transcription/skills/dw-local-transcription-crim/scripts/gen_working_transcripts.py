#!/usr/bin/env python3
"""Step 6c: one working transcript .docx per evidence item, written into that
item's own folder, plus a master media index.

    python gen_working_transcripts.py [workdir] [--index-to DIR]

Working transcripts are NOT court-reporter transcripts. They carry the
work-product header, every source file, timestamps, machine speaker labels, and
-- critically -- a coverage-gap warning naming how much recorded material is
missing. Never dress one up with numbered lines: it looks finished and is not.
"""
import os, sys, json, collections, datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

H = os.path.expanduser(sys.argv[1] if len(sys.argv) > 1 and not sys.argv[1].startswith("--") else "~/dw-asr")
B = os.path.join(H, "batch")
IDX_DIR = None
if "--index-to" in sys.argv: IDX_DIR = sys.argv[sys.argv.index("--index-to") + 1]

CAPTION = os.environ.get("DW_CAPTION",
    "State of Louisiana v. <CLIENT> · No. <DOCKET> · <COURT>, <PARISH> Parish")
fmt = lambda t: "%d:%02d:%02d" % (int(t)//3600, (int(t)%3600)//60, int(t)%60)

def item_of(p):
    parts = p.split("/")
    return next((q for q in reversed(parts[:-1]) if q.startswith("#")), parts[-2])

def para(doc, txt, sz=10, bold=False, ital=False, align=None, indent=None, mono=False):
    p = doc.add_paragraph(); r = p.add_run(txt)
    r.font.size = Pt(sz); r.bold = bold; r.italic = ital
    r.font.name = "Courier New" if mono else "Times New Roman"
    if align: p.alignment = align
    if indent: p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(2); return p

def work_header(section):
    hp = section.header.paragraphs[0]
    hp.text = ("ATTORNEY WORK PRODUCT — PRIVILEGED AND CONFIDENTIAL  ·  "
               "PREPARED IN ANTICIPATION OF LITIGATION")
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in hp.runs: r.font.size = Pt(7.5); r.bold = True

man = [json.loads(l) for l in open(os.path.join(H, "batch_manifest.jsonl"))]
groups = collections.defaultdict(list)
for m in man: groups[item_of(m["path"])].append(m)

index_rows, written = [], []
for item, ms in sorted(groups.items()):
    ms = sorted(ms, key=lambda m: m["path"])
    ok = [m for m in ms if m.get("status") == "OK" and os.path.exists(f"{B}/{m['id']}.json")]
    if not ok: continue
    doc = Document(); s = doc.sections[0]
    s.left_margin = s.right_margin = Inches(0.9); s.top_margin = Inches(0.8)
    work_header(s)
    para(doc, CAPTION, 9, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, item, 15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "WORKING TRANSCRIPT — machine transcription, not hand-verified",
         10, ital=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    probed = sum(m["dur"] for m in ms); used = sum(m.get("extracted_dur", 0) for m in ok)
    gaps = [m for m in ms if m.get("short") or m.get("status") != "OK"]
    para(doc, f"Source files: {len(ms)}   ·   Media runtime: {fmt(probed)}   ·   "
              f"Audio transcribed: {fmt(used)}   ·   Generated {datetime.date.today()}",
         9, align=WD_ALIGN_PARAGRAPH.CENTER)
    if gaps:
        para(doc, f"COVERAGE GAP — {len(gaps)} source file(s) in this item could not be read past "
                  f"the first 8 MB and are NOT transcribed below. Approximately {fmt(probed-used)} "
                  f"of recorded material is missing. See the media integrity audit.",
             9, bold=True, ital=True)
    para(doc, "Produced by Whisper large-v3 running locally on firm hardware, with pyannote speaker "
              "diarization where applicable. NOT verified against the recording by a human listener. "
              "Speaker labels are machine-assigned. Do not quote in a filing without checking the audio.",
         8, ital=True)
    doc.add_paragraph()
    for m in ok:
        w = json.load(open(f"{B}/{m['id']}.json"))
        tj = f"{B}/{m['id']}.turns.json"
        turns = json.load(open(tj)) if os.path.exists(tj) else None
        para(doc, os.path.basename(m["path"]), 9, bold=True)
        para(doc, f"   duration {fmt(m.get('extracted_dur', 0))}" +
                  (f"   ·   {m['speakers']} distinct voices" if m.get("speakers") else ""), 8, ital=True)
        cur = None
        for seg in w["segments"]:
            t = seg["text"].strip()
            if not t: continue
            lab = ""
            if turns:
                sc = collections.Counter()
                for x in turns:
                    o = min(seg["end"], x["end"]) - max(seg["start"], x["start"])
                    if o > 0: sc[x["spk"]] += o
                if sc:
                    top, tv = sc.most_common()[0]
                    lab = top if tv / sum(sc.values()) >= 0.7 else "?"   # margin, else unknown
            if lab and lab != cur:
                para(doc, f"[{lab}]", 8, bold=True); cur = lab
            p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Inches(0.35)
            r = p.add_run(f"{fmt(seg['start'])}  "); r.font.size = Pt(8)
            r.font.name = "Courier New"; r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            r2 = p.add_run(t); r2.font.size = Pt(10); r2.font.name = "Times New Roman"
        doc.add_paragraph()
    out = os.path.join(os.path.dirname(ok[0]["path"]),
                       f"WORKING TRANSCRIPT — {item[:60]} ({datetime.date.today()}).docx")
    try: doc.save(out)
    except Exception:
        out = os.path.join(H, "out_docx", f"{item[:60]}.docx")
        os.makedirs(os.path.dirname(out), exist_ok=True); doc.save(out)
    written.append(out)
    index_rows.append({"item": item, "files": len(ms), "probed": probed, "usable": used,
                       "trunc": len(gaps), "speakers": max([m.get("speakers") or 0 for m in ok] + [0]),
                       "iv": any(m.get("iv") for m in ms)})

json.dump(index_rows, open(os.path.join(H, "index_rows.json"), "w"))
print(f"working transcripts written: {len(written)}")

# ---- master index (landscape) --------------------------------------------
doc = Document(); s = doc.sections[0]
s.left_margin = s.right_margin = Inches(0.6); s.top_margin = Inches(0.7)
s.page_width, s.page_height = s.page_height, s.page_width
work_header(s)
para(doc, CAPTION, 10, align=WD_ALIGN_PARAGRAPH.CENTER)
para(doc, "MASTER MEDIA TRANSCRIPTION INDEX", 16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
para(doc, f"Generated {datetime.date.today()} · All processing local to firm hardware · "
          f"No cloud transcription service used", 9, ital=True, align=WD_ALIGN_PARAGRAPH.CENTER)
tp = sum(r["probed"] for r in index_rows); tu = sum(r["usable"] for r in index_rows)
para(doc, f"{len(index_rows)} evidence items · {sum(r['files'] for r in index_rows)} media files · "
          f"{fmt(tp)} recorded · {fmt(tu)} transcribed · {fmt(tp-tu)} UNRECOVERABLE",
     10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
index_rows.sort(key=lambda r: (not r["iv"], -r["probed"]))
t = doc.add_table(rows=1, cols=7); t.style = "Table Grid"
for j, htxt in enumerate(["Evidence item", "Type", "Files", "Recorded", "Transcribed", "Voices", "Coverage"]):
    c = t.rows[0].cells[j]; c.text = htxt
    for p in c.paragraphs:
        for r in p.runs: r.bold = True; r.font.size = Pt(8.5)
for r in index_rows:
    pct = 100 * r["usable"] / r["probed"] if r["probed"] else 0
    cov = "COMPLETE" if pct > 95 else (f"PARTIAL {pct:.0f}%" if pct > 10 else "NO AUDIO RECOVERED")
    cells = t.add_row().cells
    for j, v in enumerate([r["item"][:56], "Interview" if r["iv"] else "Other", str(r["files"]),
                           fmt(r["probed"]), fmt(r["usable"]), str(r["speakers"] or "-"), cov]):
        cells[j].text = v
        for p in cells[j].paragraphs:
            for run in p.runs:
                run.font.size = Pt(8.5); run.font.name = "Times New Roman"
                if j == 6 and cov.startswith("NO"): run.bold = True
out = os.path.join(IDX_DIR or H, f"MASTER MEDIA TRANSCRIPTION INDEX ({datetime.date.today()}).docx")
doc.save(out); print("index:", out)
