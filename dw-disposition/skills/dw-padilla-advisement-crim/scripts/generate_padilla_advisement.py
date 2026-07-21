#!/usr/bin/env python3
"""
Generate a signable D&W Padilla immigration-consequences advisement (.docx).

Usage:
  python3 generate_padilla_advisement.py --params case_params.json \
      --language-file ../assets/advisement_language.json --out "OUTPUT.docx"

case_params.json schema:
{
  "client_name": "Elvin Ulloa Meza",
  "client_dob": "09/06/1996",
  "docket": "18033-25",
  "court": "14th Judicial District Court, Division H, Parish of Calcasieu",
  "judge": "Hon. Kendrick Guidry",
  "attorney": "Christopher Washington, Daniels & Washington, LLC",
  "date": "____________, 2026",
  "charges_text": "Ct. 1 — La. R.S. 14:42 First Degree Rape;  Ct. 2 — La. R.S. 14:43.1(C)(2) Sexual Battery of a victim under 13",
  "charge_classification": {
      "en": "The offenses charged in this case ... are classified under federal immigration law as \"aggravated felonies\" ... crimes involving moral turpitude ... and crimes of child abuse.",
      "es": "Los delitos imputados en este caso ... están clasificados ... como \"delitos agravados\" ... \"delitos de vileza moral\" ... y \"delitos de maltrato infantil\"."
  },
  "conditional_flags": ["sex_offense_registration"],       # any of: sex_offense_registration, controlled_substance, firearm, domestic_violence
  "include_deportation_caution": true,
  "sentence_consequence_en": "a long Louisiana prison sentence (Count 1 carries mandatory life without parole)",
  "sentence_consequence_es": "una larga condena de prisión en Luisiana (el Cargo 1 conlleva cadena perpetua obligatoria sin libertad condicional)",
  "language": "bilingual",                                  # "bilingual" or "en"
  "interpreter_language": "Spanish"
}
"""
import argparse, json, os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
RED = RGBColor(0xB0, 0x00, 0x20)
GREY_FILL = "F2F2F2"
NAVY_FILL = "1F3A5F"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def set_cell_bg(cell, hexfill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexfill)
    tcPr.append(shd)

def run(p, text, bold=False, italic=False, size=11, color=None):
    r = p.add_run(text); r.bold = bold; r.italic = italic
    r.font.size = Pt(size); r.font.name = 'Calibri'
    if color is not None: r.font.color.rgb = color
    return r

def band(doc, en, es, bilingual):
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    c = t.cell(0, 0); set_cell_bg(c, NAVY_FILL)
    c.paragraphs[0].text = ''
    run(c.paragraphs[0], en, bold=True, size=11, color=WHITE)
    if bilingual:
        p2 = c.add_paragraph(); run(p2, es, bold=True, italic=True, size=10, color=WHITE)
    return t

def numbered_bi_table(doc, rows, bilingual, start=1):
    """rows: list of dict {en, es}. Two-column EN/ES table with numbers, or single col if en-only."""
    ncols = 2 if bilingual else 1
    tbl = doc.add_table(rows=1, cols=ncols); tbl.style = 'Table Grid'; tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    set_cell_bg(hdr[0], NAVY_FILL); hdr[0].paragraphs[0].text=''
    run(hdr[0].paragraphs[0], 'ENGLISH', bold=True, color=WHITE)
    if bilingual:
        set_cell_bg(hdr[1], NAVY_FILL); hdr[1].paragraphs[0].text=''
        run(hdr[1].paragraphs[0], 'ESPAÑOL', bold=True, color=WHITE)
    n = start
    for row in rows:
        cells = tbl.add_row().cells
        cells[0].paragraphs[0].text=''
        run(cells[0].paragraphs[0], f"{n}. {row['en']}")
        if bilingual:
            set_cell_bg(cells[1], GREY_FILL); cells[1].paragraphs[0].text=''
            run(cells[1].paragraphs[0], f"{n}. {row['es']}", italic=True)
        n += 1
    _equal_cols(tbl, ncols)
    return n

def _equal_cols(tbl, ncols):
    w = Inches(6.5/ncols)
    for r in tbl.rows:
        for c in r.cells:
            c.width = w

def para(doc, en, es, bilingual, italic_es=True, after=6):
    p = doc.add_paragraph(); run(p, en); p.paragraph_format.space_after = Pt(after)
    if bilingual:
        p2 = doc.add_paragraph(); run(p2, es, italic=italic_es); p2.paragraph_format.space_after = Pt(after+6)

def sig_block(doc, label_en, label_es, bilingual):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10)
    run(p, "_________________________________________     Date/Fecha: __________________")
    p2 = doc.add_paragraph()
    run(p2, label_en, bold=True)
    if bilingual: run(p2, "  /  "); run(p2, label_es, italic=True)
    p2.paragraph_format.space_after = Pt(8)

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('--params', required=True)
    ap.add_argument('--language-file', required=True)
    ap.add_argument('--out', required=True)
    a = ap.parse_args()
    P = json.load(open(a.params, encoding='utf-8'))
    L = json.load(open(a.language_file, encoding='utf-8'))
    bilingual = P.get('language', 'bilingual') == 'bilingual'

    doc = Document()
    sec = doc.sections[0]
    for m in ('top_margin','bottom_margin','left_margin','right_margin'): setattr(sec, m, Inches(0.75))

    # Header
    h = doc.add_paragraph(); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(h, "DANIELS & WASHINGTON, LLC", bold=True, size=13, color=NAVY)
    h2 = doc.add_paragraph(); h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(h2, "Criminal Defense", size=9, color=NAVY)
    h3 = doc.add_paragraph(); h3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(h3, "ATTORNEY-CLIENT PRIVILEGED & CONFIDENTIAL" + ("  /  CONFIDENCIAL — PRIVILEGIO ABOGADO-CLIENTE" if bilingual else ""), bold=True, size=8, color=RED)
    t1 = doc.add_paragraph(); t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(t1, L['title']['en'], bold=True, size=12)
    if bilingual:
        t2 = doc.add_paragraph(); t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run(t2, L['title']['es'], bold=True, italic=True, size=10)

    # Case identifier table
    idrows = [
        ("Client", "Cliente", P.get('client_name','') + (f"  (DOB {P['client_dob']})" if P.get('client_dob') else "")),
        ("Docket / Court", "Expediente / Tribunal", f"No. {P.get('docket','')} — {P.get('court','')}"),
        ("Judge", "Juez", P.get('judge','')),
        ("Charges", "Cargos", P.get('charges_text','')),
        ("Attorney", "Abogado", P.get('attorney','')),
        ("Date", "Fecha", P.get('date','____________')),
    ]
    it = doc.add_table(rows=0, cols=2); it.style='Table Grid'
    for en,es,val in idrows:
        cells = it.add_row().cells; set_cell_bg(cells[0], GREY_FILL)
        cells[0].paragraphs[0].text=''
        run(cells[0].paragraphs[0], en, bold=True, size=10)
        if bilingual: run(cells[0].paragraphs[0], "  /  "); run(cells[0].paragraphs[0], es, italic=True, size=9)
        cells[1].paragraphs[0].text=''; run(cells[1].paragraphs[0], val, size=10)
    it.columns[0].width = Inches(2.4); it.columns[1].width = Inches(4.1)
    for r in it.rows:
        r.cells[0].width = Inches(2.4); r.cells[1].width = Inches(4.1)
    doc.add_paragraph()

    # Purpose
    band(doc, "Purpose of This Advisement", "Propósito de esta advertencia", bilingual)
    para(doc, L['purpose']['en'], L['purpose']['es'], bilingual)

    # Build the main numbered points
    band(doc, "Immigration Consequences You Must Understand", "Consecuencias migratorias que usted debe entender", bilingual)
    points = []
    for pt in L['universal_points']:
        if pt['en'] == "__CHARGE_CLASSIFICATION__":
            cc = P.get('charge_classification', {})
            points.append({'en': cc.get('en','[CHARGE CLASSIFICATION — attorney to complete]'),
                           'es': cc.get('es','[CLASIFICACIÓN DE CARGOS — a completar por el abogado]')})
        else:
            points.append(pt)
    # insert conditional points (after point 7 = last universal) in a stable order
    for flag in ['sex_offense_registration','controlled_substance','firearm','domestic_violence']:
        if flag in P.get('conditional_flags', []):
            points.append(L['conditional_points'][flag])
    nxt = numbered_bi_table(doc, points, bilingual, start=1)
    doc.add_paragraph()

    # Deportation caution (optional)
    if P.get('include_deportation_caution'):
        dc = L['deportation_caution']
        band(doc, dc['heading']['en'], dc['heading']['es'], bilingual)
        dpoints = []
        se = P.get('sentence_consequence_en','a criminal sentence')
        ses = P.get('sentence_consequence_es','una condena penal')
        for pt in dc['points']:
            dpoints.append({'en': pt['en'].replace('__SENTENCE_CONSEQUENCE__', se),
                            'es': pt['es'].replace('__SENTENCE_CONSEQUENCE_ES__', ses)})
        nxt = numbered_bi_table(doc, dpoints, bilingual, start=nxt)
        doc.add_paragraph()

    # Rights & recommendation
    band(doc, "Your Rights and Our Recommendation", "Sus derechos y nuestra recomendación", bilingual)
    numbered_bi_table(doc, L['rights_points'], bilingual, start=nxt)
    doc.add_paragraph()

    # Certifications
    for key in ['client','attorney','interpreter']:
        c = L['certifications'][key]
        band(doc, c['heading']['en'], c['heading']['es'], bilingual)
        para(doc, c['en'], c['es'], bilingual)
        name_en = c['sig_en']
        if key == 'client' and P.get('client_name'):
            name_en = f"{c['sig_en']} — {P['client_name']}"
        if key == 'attorney' and P.get('attorney'):
            name_en = f"{c['sig_en']} — {P['attorney']}"
        sig_block(doc, name_en, c['sig_es'], bilingual)

    # Footer note
    fp = doc.add_paragraph(); fp.paragraph_format.space_before = Pt(12)
    run(fp, L['footer']['en'], italic=True, size=8, color=RGBColor(0x7F,0x7F,0x7F))

    doc.save(a.out)
    print("WROTE", a.out)

if __name__ == '__main__':
    main()
